import os
import tempfile
import base64
import json
import httpx
import re
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
from typing import Dict, Any, Optional, List
from routes.auth import get_current_approved_user
from config import settings
from database import log_audit_action
from services.task_service import task_queue

router = APIRouter(prefix="/reports", tags=["reports"])

class ReportRequest(BaseModel):
    bearing_name: str
    primary_diagnosis: str
    confidence_score: int
    recommendations: str
    critical_speeds: Dict[str, Any]
    telemetry_summary: Dict[str, Any]

@router.post("/generate")
async def generate_report(
    payload: ReportRequest,
    current_user: dict = Depends(get_current_approved_user)
):
    sub_status = current_user.get("subscription_status", "free-tier")
    if sub_status != "premium":
        gen_count = current_user.get("report_generation_count", 0)
        if gen_count >= 3:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Free-tier report generation limit of 3 reports exhausted. Upgrade to a Premium subscription for unlimited reports."
            )
        
        # Increment generation counter in profiles and user_metadata
        new_count = gen_count + 1
        
        # 1. Update user_metadata (Auth System)
        try:
            supabase.auth.admin.update_user_by_id(
                current_user["id"],
                {"user_metadata": {"report_generation_count": new_count}}
            )
        except Exception as auth_err:
            print(f"Failed to increment report_generation_count in user_metadata: {auth_err}")
            
        # 2. Update profiles table (DB System, ignore error if column is missing)
        try:
            supabase.table("profiles").update({"report_generation_count": new_count}).eq("id", current_user["id"]).execute()
        except Exception as db_err:
            print(f"Failed to update profiles table column (expected if column is missing): {db_err}")

    if not settings.GEMINI_API_KEY:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Gemini API Key is not configured on the backend server."
        )

    # Build prompt
    prompt = f"""
    You are an expert rotating machinery engineer and senior vibration analyst.
    Please generate a professional turbomachinery engineering maintenance report based on the following diagnostic metadata:

    - Probe Location: {payload.bearing_name}
    - Primary Diagnosed Issue: {payload.primary_diagnosis}
    - Diagnostic Confidence Score: {payload.confidence_score}%
    - Standard Heuristic Recommendations: {payload.recommendations}
    
    Identified Critical Speeds (Resonance):
    {payload.critical_speeds}

    Telemetry and Operation Bounds:
    {payload.telemetry_summary}

    Please write a comprehensive, highly technical machinery health report. The report must contain the following structured sections:
    1. EXECUTIVE SUMMARY
       - Brief status overview, critical alerts, and core diagnostic conclusion.
    2. SHAFTPATH & ORBIT VIBRATION ANALYSIS
       - Detailed interpretation of the proximity probe orbits, centerline behavior, and dynamic clearance headroom.
    3. DETECTED FAULT MECHANISMS
       - Physical engineering explanation of the diagnosed issues (e.g. unbalance, misalignment, rub, oil whirl) and their root causes.
    4. MAINTENANCE ACTION PLAN & CORRECTIVE RECOMMENDATIONS
       - Specific actionable steps (e.g. dynamic balancing, realignment, check bearing clearances, rotor shaft runout tests).
    5. RECOMMENDED MONITORING & OPERATION SCHEDULE
       - Immediate, short-term, and long-term operating intervals and alarm parameters.

    Format the report entirely in clean, readable Markdown. Do not include raw HTML code. Be thorough, technical, and professional in your writing.
    """

    url = f"https://generativelanguage.googleapis.com/v1/models/gemini-2.5-flash:streamGenerateContent?key={settings.GEMINI_API_KEY}"
    
    headers = {
        "Content-Type": "application/json"
    }
    
    data = {
        "contents": [
            {
                "parts": [
                    {
                        "text": prompt
                    }
                ]
            }
        ]
    }
    
    log_audit_action(current_user["id"], "GENERATE_AI_REPORT", {
        "bearing_name": payload.bearing_name,
        "primary_diagnosis": payload.primary_diagnosis,
        "confidence_score": payload.confidence_score
    })
    
    async def stream_generator():
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                async with client.stream("POST", url, headers=headers, json=data) as response:
                    if response.status_code != 200:
                        yield f"data: {json.dumps({'error': f'Gemini API returned status {response.status_code}'})}\n\n"
                        return
                    
                    buffer = ""
                    async for line in response.aiter_lines():
                        if not line:
                            continue
                        buffer += line
                        clean_buffer = buffer.strip().lstrip(',[').rstrip(',]')
                        try:
                            chunk_data = json.loads(clean_buffer)
                            candidates = chunk_data.get("candidates", [])
                            if candidates:
                                text_val = candidates[0].get("content", {}).get("parts", [{}])[0].get("text", "")
                                if text_val:
                                    yield f"data: {json.dumps({'text': text_val})}\n\n"
                            buffer = ""
                        except Exception:
                            # Incomplete JSON, accumulate in buffer
                            pass
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(stream_generator(), media_type="text/event-stream")

class DocxReportRequest(BaseModel):
    bearing_name: str
    report_text: str
    images: Optional[List[str]] = []

@router.post("/download_docx")
async def download_docx(
    payload: DocxReportRequest,
    current_user: dict = Depends(get_current_approved_user)
):
    if current_user.get("subscription_status", "free-tier") != "premium":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Upgrade to a Premium subscription is required to export reports as Word Documents."
        )

    try:
        temp_dir = tempfile.gettempdir()
        safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', payload.bearing_name)
        temp_file_path = os.path.join(temp_dir, f"RotorDyn_Report_{safe_name}.docx")
        
        doc = docx.Document()
        
        # Margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        # Global Font
        style_normal = doc.styles['Normal']
        font = style_normal.font
        font.name = 'Arial'
        font.size = Pt(10.5)
        font.color.rgb = docx.shared.RGBColor(15, 23, 42)
        
        # Header
        title = doc.add_paragraph()
        title_run = title.add_run("ROTORDYN.AI")
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = docx.shared.RGBColor(2, 132, 199)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        sub_run = subtitle.add_run("Machinery Diagnostics & Rotor Dynamic Analysis Report")
        sub_run.bold = True
        sub_run.font.size = Pt(13)
        sub_run.font.color.rgb = docx.shared.RGBColor(71, 85, 105)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        divider = doc.add_paragraph()
        div_run = divider.add_run("__________________________________________________________________")
        div_run.font.color.rgb = docx.shared.RGBColor(203, 213, 225)
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Helper to parse **bold** text in markdown paragraphs
        def parse_inline_bolding(paragraph, text):
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    paragraph.add_run(part)
        
        # Parse and write report markdown
        lines = payload.report_text.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue
            
            # Match headers
            if stripped.startswith('### '):
                h = doc.add_heading(stripped[4:], level=3)
                if len(h.runs) > 0:
                    h.runs[0].font.color.rgb = docx.shared.RGBColor(71, 85, 105)
            elif stripped.startswith('## '):
                h = doc.add_heading(stripped[3:], level=2)
                if len(h.runs) > 0:
                    h.runs[0].font.color.rgb = docx.shared.RGBColor(15, 23, 42)
            elif stripped.startswith('# '):
                h = doc.add_heading(stripped[2:], level=1)
                if len(h.runs) > 0:
                    h.runs[0].font.color.rgb = docx.shared.RGBColor(2, 132, 199)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                text = stripped[2:]
                p = doc.add_paragraph(style='List Bullet')
                parse_inline_bolding(p, text)
            else:
                p = doc.add_paragraph()
                parse_inline_bolding(p, stripped)

        # Section 6: System Evidence & Diagnostic Plots
        if payload.images:
            doc.add_page_break()
            doc.add_heading("System Evidence & Diagnostic Plots", level=1)
            doc.add_paragraph("The following plots are captured directly from the active diagnostic workspace environment as empirical evidence of the analyzed machinery behavior:")
            
            for idx, img_base64 in enumerate(payload.images):
                try:
                    if ',' in img_base64:
                        img_base64 = img_base64.split(',')[1]
                    
                    # Decode image and write to temp file
                    img_data = base64.b64decode(img_base64)
                    img_temp_path = os.path.join(temp_dir, f"evidence_plot_{idx}.png")
                    with open(img_temp_path, "wb") as fh:
                        fh.write(img_data)
                    
                    # Add picture
                    doc.add_paragraph(f"\nEvidence Plot #{idx+1}:", style='Normal').runs[0].font.bold = True
                    doc.add_picture(img_temp_path, width=Inches(5.8))
                    
                    # Cleanup temp image file
                    try:
                        os.remove(img_temp_path)
                    except:
                        pass
                except Exception as e_img:
                    print("Error adding image to docx:", e_img)
                    
        doc.save(temp_file_path)
        
        # Log audit
        log_audit_action(current_user["id"], "EXPORT_DOCX_REPORT", {
            "bearing_name": payload.bearing_name,
            "has_images": len(payload.images) > 0
        })
        
        return FileResponse(
            path=temp_file_path,
            filename=f"RotorDyn_Report_{safe_name}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate Word document: {str(e)}"
        )

async def compile_gemini_report_task(payload_dict: dict, user_id: str):
    """Background worker task compilation handler that calls Gemini API non-blocking."""
    api_key = settings.GEMINI_API_KEY
    if not api_key:
        raise ValueError("Gemini API Key is not configured.")
        
    prompt = f"""
    You are an expert rotating machinery engineer and senior vibration analyst.
    Please generate a professional turbomachinery engineering maintenance report based on the following diagnostic metadata:

    - Probe Location: {payload_dict.get('bearing_name')}
    - Primary Diagnosed Issue: {payload_dict.get('primary_diagnosis')}
    - Diagnostic Confidence Score: {payload_dict.get('confidence_score')}%
    - Standard Heuristic Recommendations: {payload_dict.get('recommendations')}
    
    Identified Critical Speeds (Resonance):
    {payload_dict.get('critical_speeds')}

    Telemetry and Operation Bounds:
    {payload_dict.get('telemetry_summary')}

    Please write a comprehensive, highly technical machinery health report.
    """
    
    url = f"https://generativelanguage.googleapis.com/v1/models/gemini-2.5-flash:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {
        "contents": [{"parts": [{"text": prompt}]}]
    }
    
    async with httpx.AsyncClient(timeout=60.0) as client:
        res = await client.post(url, headers=headers, json=data)
        if res.status_code != 200:
            raise Exception(f"Gemini API returned error code {res.status_code}")
        
        result_json = res.json()
        candidates = result_json.get("candidates", [])
        if not candidates:
            raise Exception("No text candidates returned from Gemini model response.")
            
        text_val = candidates[0]["content"]["parts"][0]["text"]
        
        log_audit_action(user_id, "GENERATE_AI_REPORT_BACKGROUND", {
            "bearing_name": payload_dict.get('bearing_name'),
            "primary_diagnosis": payload_dict.get('primary_diagnosis')
        })
        return text_val

@router.post("/generate_async", status_code=status.HTTP_202_ACCEPTED)
async def generate_report_async(
    payload: ReportRequest,
    current_user: dict = Depends(get_current_approved_user)
):
    """Enqueues generative report compilation dynamically using swappable task queue."""
    sub_status = current_user.get("subscription_status", "free-tier")
    if sub_status != "premium":
        gen_count = current_user.get("report_generation_count", 0)
        if gen_count >= 3:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Free-tier report generation limit of 3 reports exhausted. Upgrade to a Premium subscription for unlimited reports."
            )
        
        new_count = gen_count + 1
        # Update user metadata & profiles count
        try:
            supabase.auth.admin.update_user_by_id(
                current_user["id"],
                {"user_metadata": {"report_generation_count": new_count}}
            )
            supabase.table("profiles").update({"report_generation_count": new_count}).eq("id", current_user["id"]).execute()
        except Exception:
            pass

    job_id = await task_queue.enqueue(
        compile_gemini_report_task,
        payload.dict(),
        current_user["id"]
    )
    return {"status": "queued", "job_id": job_id}

@router.get("/status/{job_id}")
async def get_report_job_status(
    job_id: str,
    current_user: dict = Depends(get_current_approved_user)
):
    """Retrieve state/result of enqueued background report compilation jobs."""
    status_info = await task_queue.get_status(job_id)
    if not status_info:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job ID not found or expired."
        )
    return status_info

